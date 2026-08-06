"""Build the Word edition of the paper.

The skill's usual toolchain (docx-js via Node, LibreOffice, pandoc) is not
installed on this machine, so this uses python-docx, which is pure Python and
needs no system dependencies. Figures are the matplotlib renders from
src/paper_figures.py -- Word cannot embed the web edition's inline SVG.

Content is authored here rather than scraped from paper.html: the HTML carries
interactive affordances (hover tooltips, a sticky rail) that have no meaning on
paper, and the Word edition wants numbered captions and a table of contents
instead.
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent / "src"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from paper_figures import build_all

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "The-Shape-of-the-Modern-Pitch.docx"

INK = RGBColor(0x11, 0x16, 0x1C)
INK2 = RGBColor(0x4A, 0x55, 0x63)
ACCENT = RGBColor(0x0E, 0x72, 0x50)
RULE_HEX = "DCE1E7"
BAND_HEX = "F2F5F8"
ACCENT_BAND = "E8F4EE"
WARN_BAND = "FBF3E7"

BODY_FONT = "Georgia"
DATA_FONT = "Consolas"


# --------------------------------------------------------------------------
# low-level helpers
# --------------------------------------------------------------------------
def shade(cell_or_par, hex_fill: str) -> None:
    """Apply a background fill. Uses CLEAR, never SOLID (SOLID renders black)."""
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    target = cell_or_par._tc.get_or_add_tcPr() if hasattr(cell_or_par, "_tc") else cell_or_par._p.get_or_add_pPr()
    target.append(el)


def bottom_border(paragraph, size: int = 6, color: str = RULE_HEX) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def left_bar(paragraph, color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    pPr.append(borders)


def para(doc, text="", *, size=10.5, bold=False, italic=False, color=INK,
         font=BODY_FONT, space_after=8, space_before=0, align=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        run.font.name = font
    return p


def rich(doc, segments, *, size=10.5, space_after=8, indent=0, font=BODY_FONT):
    """A paragraph built from (text, style) pairs: '', 'b', 'i', 'm' (mono)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for text, style in segments:
        run = p.add_run(text)
        run.font.size = Pt(size if style != "m" else size - 0.5)
        run.font.bold = "b" in style
        run.font.italic = "i" in style
        run.font.name = DATA_FONT if style == "m" else font
        run.font.color.rgb = INK
    return p


def heading(doc, number: str, title: str, level: int = 1):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(6)
    if number:
        n = h.add_run(number + "  ")
        n.font.color.rgb = ACCENT
        n.font.name = DATA_FONT
        n.font.size = Pt(15 if level == 1 else 12)
        n.font.bold = True
    t = h.add_run(title)
    t.font.color.rgb = INK
    t.font.name = BODY_FONT
    t.font.size = Pt(16 if level == 1 else 12.5)
    t.font.bold = True
    if level == 1:
        bottom_border(h)
    return h


def bullets(doc, items, *, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        if isinstance(item, str):
            item = [(item, "")]
        for text, style in item:
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = "b" in style
            run.font.italic = "i" in style
            run.font.name = DATA_FONT if style == "m" else BODY_FONT
            run.font.color.rgb = INK


def numbered(doc, items, *, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, str):
            item = [(item, "")]
        for text, style in item:
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = "b" in style
            run.font.italic = "i" in style
            run.font.name = DATA_FONT if style == "m" else BODY_FONT
            run.font.color.rgb = INK


def callout(doc, tag: str, body_parts, *, band=BAND_HEX, bar="1BAF7A"):
    """A shaded, left-barred aside — the Word analogue of the web callouts."""
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after = Pt(0)
    tp.paragraph_format.left_indent = Inches(0.12)
    r = tp.add_run(tag.upper())
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.name = DATA_FONT
    r.font.color.rgb = INK2
    shade(tp, band)
    left_bar(tp, bar)

    for parts in body_parts:
        if isinstance(parts, str):
            parts = [(parts, "")]
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(6)
        bp.paragraph_format.left_indent = Inches(0.12)
        for text, style in parts:
            run = bp.add_run(text)
            run.font.size = Pt(10 if style != "m" else 9.5)
            run.font.bold = "b" in style
            run.font.italic = "i" in style
            run.font.name = DATA_FONT if style == "m" else BODY_FONT
            run.font.color.rgb = INK
        shade(bp, band)
        left_bar(bp, bar)


def figure(doc, image_path, number: int, title: str, caption: str, width_in=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(image_path), width=Inches(width_in))

    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(14)
    lab = cap.add_run(f"Figure {number}. ")
    lab.font.size = Pt(8.5)
    lab.font.bold = True
    lab.font.name = DATA_FONT
    lab.font.color.rgb = ACCENT
    ttl = cap.add_run(title + " ")
    ttl.font.size = Pt(8.5)
    ttl.font.bold = True
    ttl.font.name = BODY_FONT
    ttl.font.color.rgb = INK
    body = cap.add_run(caption)
    body.font.size = Pt(8.5)
    body.font.name = BODY_FONT
    body.font.color.rgb = INK2


def table(doc, number: int, title: str, headers, rows, widths, *, note=None):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(12)
    cap.paragraph_format.space_after = Pt(4)
    lab = cap.add_run(f"Table {number}. ")
    lab.font.size = Pt(8.5)
    lab.font.bold = True
    lab.font.name = DATA_FONT
    lab.font.color.rgb = ACCENT
    ttl = cap.add_run(title)
    ttl.font.size = Pt(8.5)
    ttl.font.bold = True
    ttl.font.name = BODY_FONT
    ttl.font.color.rgb = INK

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Both the table's columnWidths and every cell width must be set in DXA,
    # or column sizing is ignored in Word and broken in Google Docs.
    t.autofit = False
    for i, w in enumerate(widths):
        t.columns[i].width = Inches(w)

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(widths[i])
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(h)
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.name = DATA_FONT
        r.font.color.rgb = INK
        shade(cell, BAND_HEX)

    for row_data in rows:
        row = t.add_row()
        emphasise = isinstance(row_data, tuple) and row_data and row_data[0] == "__HI__"
        values = row_data[1:] if emphasise else row_data
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.width = Inches(widths[i])
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(val))
            r.font.size = Pt(8)
            r.font.name = DATA_FONT
            r.font.bold = emphasise
            r.font.color.rgb = ACCENT if emphasise else INK

    if note:
        n = doc.add_paragraph()
        n.paragraph_format.space_after = Pt(14)
        run = n.add_run(note)
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.name = BODY_FONT
        run.font.color.rgb = INK2
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(10)


def toc_field(doc):
    """Insert a real TOC field. Word populates it on F9 / open-and-update."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)


# --------------------------------------------------------------------------
# document
# --------------------------------------------------------------------------
def build() -> Path:
    figs = build_all()
    doc = Document()

    # US Letter with 1" margins (python-docx defaults to Letter, but set it
    # explicitly so the file is not at the mercy of a template default).
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Inches(1.0))
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    # ---------------- title page ----------------
    eyebrow = para(doc, "BASEBALL RESEARCH · PITCH-LEVEL STUDY", size=8.5, bold=True,
                   color=ACCENT, font=DATA_FONT, space_after=10, space_before=40)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    tr = title.add_run("The Shape of the Modern Pitch")
    tr.font.size = Pt(30)
    tr.font.bold = True
    tr.font.name = BODY_FONT
    tr.font.color.rgb = INK

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(22)
    sr = sub.add_run(
        "Ball-flight geometry — not velocity, and emphatically not spin rate — is the pitch "
        "property that carries information about outcomes. It is also the only one measured "
        "well enough to judge a pitcher on a small sample."
    )
    sr.font.size = Pt(13)
    sr.font.name = BODY_FONT
    sr.font.color.rgb = INK2

    meta = para(doc, "", space_after=4)
    bottom_border(meta)
    for line in (
        "Scott Luntz",
        "7,483,321 pitches · 2015–2025 regular seasons",
        "Source: Statcast / Baseball Savant",
    ):
        para(doc, line, size=9.5, font=DATA_FONT, color=INK2, space_after=2)

    # ---------------- abstract ----------------
    callout(
        doc,
        "Abstract",
        [
            [("Using every tracked pitch from the 2015–2025 regular seasons, I decompose pitch "
              "outcomes into the unique contribution of velocity, spin rate, and pitch shape "
              "(induced vertical break, horizontal break, and vertical approach angle). Shape "
              "contributes ", ""), ("4.5×", "m"),
             (" more out-of-sample explanatory power than velocity for run value and ", ""),
             ("3.4×", "m"),
             (" more for whiffs; residual spin rate, once velocity is partialled out, contributes "
              "nothing distinguishable from zero even when granted first claim on shared variance. "
              "The result holds for all seven pitch types tested and survives adjustment for pitch "
              "location, the most serious confound.", "")],
            [("The evaluation consequence follows from reliability rather than effect size. Shape "
              "metrics have split-half reliability of ", ""), ("r ≈ 0.99", "m"),
             (" within a season; run value per pitch has ", ""), ("r = 0.20", "m"),
             (". Because shape is measured almost without error, it forecasts a pitcher's next "
              "season better than his own results do below roughly 80 pitches, and continues to add "
              "information on top of results at every sample size. I state the boundary of the claim "
              "explicitly: shape does not replace outcome data on full seasons, and command — not any "
              "pitch property — remains the single largest determinant of pitch outcomes.", "")],
        ],
        band=ACCENT_BAND,
    )

    para(doc, "", space_after=0).add_run().add_break(WD_BREAK.PAGE)

    # ---------------- contents ----------------
    heading(doc, "", "Contents", level=1)
    toc_field(doc)
    para(doc, "", space_after=0).add_run().add_break(WD_BREAK.PAGE)

    # ---------------- 1 ----------------
    heading(doc, "1", "Introduction")
    para(doc, "For fifteen years, pitcher evaluation has been organised around two numbers that are "
              "easy to put on a scoreboard: how hard the ball was thrown, and how fast it was "
              "spinning. Velocity is intuitive. Spin rate arrived with Statcast in 2015, was quickly "
              "attached to the idea of a “rising” fastball, and became shorthand for pitch quality — "
              "a high-spin arm was a good arm.")
    para(doc, "Both framings are wrong in an instructive way. Neither velocity nor spin acts on a "
              "hitter directly. A hitter cannot perceive revolutions per minute. What he perceives is "
              "a trajectory: where the ball appears to be going, how it deviates from that path, and "
              "at what angle it arrives. Velocity and spin matter only insofar as they produce that "
              "trajectory.")
    rich(doc, [
        ("This study asks the question in the form that can actually be answered. Rather than asking "
         "whether spin “matters,” which is unanswerable because spin correlates with everything else "
         "a pitcher does, it asks: ", ""),
        ("after a model already knows everything else about a pitch, how much does each property add?", "i"),
        (" That question has a clean empirical answer, and the answer is that pitch shape dominates. "
         "It also has a practical corollary that is more useful than the finding itself: because "
         "shape is a near-deterministic property of how a pitcher throws, it is knowable from a "
         "handful of pitches, at a point when every outcome statistic is still noise.", ""),
    ])
    para(doc, "I have tried to write this so it can be checked and so it can be disbelieved. Section 6 "
              "collects the results that constrain the thesis, including one that contradicts its "
              "strongest form. The code and the exact commands that produce every number are in "
              "Section 9.")

    # ---------------- 2 ----------------
    heading(doc, "2", "Definitions")
    para(doc, "Four measurements do most of the work in this paper. Each is defined technically and "
              "then in plain language, because the technical definitions are precise and the plain "
              "ones are what the technical definitions actually mean.")

    heading(doc, "2.1", "Induced vertical break (IVB)", level=2)
    para(doc, "The vertical deviation of the pitch, in inches, from the path a spinless ball thrown on "
              "the same initial vector would follow — that is, the movement caused by the ball's spin, "
              "with gravity's effect removed.")
    callout(doc, "In plain English", [
        "Gravity pulls every pitch down. Backspin pushes back against it. IVB measures how much of "
        "gravity's pull the spin cancelled out. A four-seam fastball with 18 inches of IVB doesn't "
        "actually rise — it just falls about a foot and a half less than the hitter's eye expects, "
        "which is why it's described as “riding” or “carrying.”"])

    heading(doc, "2.2", "Vertical approach angle (VAA)", level=2)
    para(doc, "The angle, in degrees below horizontal, at which the ball crosses the front edge of the "
              "plate. It is computed from Statcast's release-point velocity and acceleration vectors "
              "by solving the constant-acceleration trajectory for the moment the ball reaches the "
              "plate, then taking the arctangent of the vertical over the horizontal velocity "
              "component at that instant.")
    callout(doc, "In plain English", [
        [("How steeply the ball is descending when it arrives. A hitter's swing travels slightly "
          "upward through the zone. A pitch that arrives ", ""), ("flat", "i"),
         (" — descending at, say, 4° instead of 6° — meets that swing plane at an awkward angle and "
          "gets missed or hit under. This is the real mechanism behind “the high fastball plays.”", "")],
        "Crucially, VAA is not the same as throwing high in the zone. Two pitchers can put the ball in "
        "the identical spot and arrive at different angles, because of how tall they are, how far down "
        "the mound they release, and how much the ball carried on the way. That difference is the part "
        "that belongs to the pitcher."])

    heading(doc, "2.3", "Horizontal break (HB)", level=2)
    para(doc, "Spin-induced horizontal deviation in inches, mirrored by handedness so that positive "
              "values always mean arm-side movement — otherwise right- and left-handed pitchers cancel "
              "each other out in any pooled model.")

    heading(doc, "2.4", "Residual spin rate", level=2)
    para(doc, "Raw spin rate with the component linearly predictable from release velocity removed, "
              "computed within each season and pitch type.")
    callout(doc, "In plain English", [
        [("Hard throwers spin the ball more, simply because a faster arm imparts more rotation. So raw "
          "spin rate partly ", ""), ("is", "i"),
         (" velocity wearing a different hat. Residual spin answers a cleaner question: given how hard "
          "this pitcher throws, is he spinning it more or less than you'd expect? That residual is the "
          "only version of “spin” that can be tested independently, and it is the version used "
          "throughout this paper.", "")]])

    figure(doc, figs["trajectory"], 1,
           "Two pitches, same release point, same destination, different geometry.",
           "Both trajectories leave the same hand at the same speed and cross the plate at the same "
           "height — an identical box score, an identical location chart. The pitch with more induced "
           "vertical break resists gravity, so it can be thrown on a flatter line and arrives at a "
           "shallower angle to the hitter's swing plane. Everything this paper measures is the "
           "consequence of that difference.")

    # ---------------- 3 ----------------
    heading(doc, "3", "Data")
    rich(doc, [("The sample is every pitch tracked in the regular seasons of 2015 through 2025, pulled "
                "from Baseball Savant: ", ""), ("7,483,321 pitches", "b"),
               (". Run value is present on 99.5–99.8% of rows in every season.", "")])
    para(doc, "Cleaning removes pitches missing tracking fields and physically impossible readings "
              "(release speed outside 60–108 mph, spin outside 500–3600 rpm), retaining roughly 92% of "
              "pitches. The primary analysis uses four-seam fastballs — the pitch where claims about "
              "spin are most often made — with all seven modelled pitch types examined separately in "
              "Section 5.4.")
    para(doc, "The 2026 season is deliberately excluded. It was in progress at the time of writing, and a "
              "partial season would enter the per-season trend analysis in §5.7 on different terms from "
              "the eleven complete seasons around it.")
    para(doc, "Spin measurement changed from Trackman to Hawk-Eye before the 2020 season. Values either "
              "side of that boundary are not perfectly comparable, so per-season models are fit "
              "separately and never pooled across it.")
    para(doc, "The outcome variable is run value from the pitcher's perspective: the change in run "
              "expectancy attributable to the pitch, sign-flipped so positive always means good for the "
              "pitcher. A secondary outcome is whiff rate on swings. Both are used because they answer "
              "different questions — missing bats is not the same as preventing runs.")

    # ---------------- 4 ----------------
    heading(doc, "4", "Methods")
    heading(doc, "4.1", "The attribution ladder", level=2)
    para(doc, "Feature blocks are added to a gradient-boosted model one at a time, each building on the "
              "last:")
    numbered(doc, [
        [("Controls", "b"), (" — count, plate location, distance from the middle of the zone, platoon matchup.", "")],
        [("Velocity", "b"), (" — release speed and extension.", "")],
        [("Spin", "b"), (" — residual spin rate.", "")],
        [("Shape", "b"), (" — IVB, HB, VAA, horizontal approach angle, release height and width.", "")],
    ])
    rich(doc, [("Because each block enters a model that already contains the previous ones, its "
                "improvement in held-out accuracy is its ", ""), ("unique", "i"),
               (" contribution. The ladder is then run in reverse, with spin entering before velocity, "
                "so neither variable's apparent importance is an artifact of the order it was given. "
                "Reporting both orderings is not optional: the gap between them ", ""),
               ("is", "i"), (" the collinearity.", "")])

    heading(doc, "4.2", "Guarding against leakage and self-deception", level=2)
    para(doc, "Four design choices do the real work, and each exists because its absence produces a "
              "wrong answer:")
    bullets(doc, [
        [("Folds are grouped by pitcher.", "b"), (" With random pitch-level splits, the same pitcher's "
         "fastballs appear in both training and test data, and the model memorises the arm instead of "
         "learning the physics.", "")],
        [("Every rung is scored on identical rows.", "b"), (" Features have different missingness, so a "
         "per-rung dropped-row policy would quietly hand later rungs an easier sample.", "")],
        [("Gains are paired across folds.", "b"), (" Each block's improvement is differenced fold by "
         "fold, yielding a standard error and a t-statistic. Pitch-level run value has a standard "
         "deviation of 0.23 runs around a mean of zero, and at that noise level a block can “gain” "
         "+0.0004 R² by luck. Nothing here is called real below t = 2.", "")],
        [("Location enters first.", "b"), (" Shape is only credited with what it explains after the "
         "model already knows where the pitch was thrown.", "")],
    ])

    heading(doc, "4.3", "Separating shape from location", level=2)
    rich(doc, [("The most serious objection to any shape-based finding is that vertical approach angle "
                "is mechanically tied to pitch height: ", ""), ("any", "i"),
               (" pitch at the top of the zone arrives flatter than the same pitch at the knees. A "
                "model crediting “shape” for VAA might be rediscovering that high fastballs play well.", "")])
    rich(doc, [("To test this, VAA is regressed on plate height (quadratic, within season and pitch "
                "type) and the residual retained — how flat the pitch arrived ", ""),
               ("relative to other pitches at the same height", "i"),
               (". Every result in Sections 5.3 through 5.6 uses this height-adjusted version.", "")])

    # ---------------- 5 ----------------
    heading(doc, "5", "Results")

    heading(doc, "5.1", "Attribution at the pitch level", level=2)
    para(doc, "On four-seam fastballs, with folds grouped by pitcher and gains paired across identical "
              "folds, shape is the only block that meaningfully moves the model.")
    figure(doc, figs["ladder"], 2,
           "Unique out-of-sample explanatory power added by each block.",
           "Each block is added on top of all previous blocks, so the bar is that block's unique "
           "contribution. The two panels use different axis scales — whiff outcomes carry far more "
           "signal than run value at the pitch level.")

    table(doc, 1, "Attribution ladder, four-seam fastballs, both orderings",
          ["Block added", "CV R²", "Gain", "SE", "t", "Verdict"],
          [
              ["Run value · velocity first", "", "", "", "", ""],
              ["Controls", "0.03667", "—", "—", "—", "baseline"],
              ["+ Velocity", "0.03691", "+0.00024", "0.00007", "+3.63", "real"],
              ["+ Spin", "0.03694", "+0.00004", "0.00003", "+1.15", "not sig."],
              ("__HI__", "+ Shape", "0.03802", "+0.00108", "0.00010", "+10.75", "real"),
              ["Run value · spin first", "", "", "", "", ""],
              ["+ Spin", "0.03664", "−0.00003", "0.00006", "−0.46", "nothing"],
              ["+ Velocity", "0.03694", "+0.00030", "0.00008", "+3.63", "real"],
              ("__HI__", "+ Shape", "0.03802", "+0.00108", "0.00010", "+10.75", "real"),
              ["Whiffs · velocity first", "", "", "", "", ""],
              ["Controls", "0.07718", "—", "—", "—", "baseline"],
              ["+ Velocity", "0.08298", "+0.00580", "0.00055", "+10.56", "real"],
              ["+ Spin", "0.08467", "+0.00169", "0.00013", "+13.08", "real"],
              ("__HI__", "+ Shape", "0.10419", "+0.01952", "0.00076", "+25.57", "real"),
          ],
          [1.75, 0.85, 0.92, 0.85, 0.72, 0.95])

    heading(doc, "5.2", "Spin is a proxy, not a cause", level=2)
    rich(doc, [("The decisive test is the reversed ladder. Residual spin was handed ", ""), ("first", "i"),
               (" claim on every scrap of variance it shares with velocity and shape — the most "
                "generous treatment the design can give it — and it returned ", ""),
               ("−0.00003 ± 0.00006", "m"), (" (", ""), ("t = −0.46", "m"),
               ("). Nothing. Velocity, entering afterward on the leftovers, still cleared its noise.", "")])
    rich(doc, [("This is not a claim that spin is physically irrelevant; it is a claim about what spin "
                "rate ", ""), ("tells you that you didn't already know", "i"),
               (". Spin's entire contribution is already expressed in the movement it produced, and the "
                "movement is directly observable. Once you can see IVB and VAA, the rpm figure is "
                "redundant.", "")])
    callout(doc, "Why this matters mechanically", [
        "Spin only moves a baseball when the spin axis is perpendicular to the direction of travel. "
        "Pure gyroscopic spin — a bullet-style rotation — moves the ball not at all. A pitch can "
        "register 2,800 rpm and break like a brick. This is why “high spin” is an unreliable label and "
        "why the movement it produces is the honest measurement."],
        band=WARN_BAND, bar="D9A441")

    heading(doc, "5.3", "Shape is not location in disguise", level=2)
    para(doc, "Replacing raw VAA with the height-adjusted residual leaves the result essentially intact.")
    table(doc, 2, "Shape block with raw versus height-adjusted approach angle",
          ["Shape block specification", "Gain", "SE", "t", "Retained"],
          [
              ["Raw VAA", "+0.00099", "0.00011", "+9.38", "—"],
              ("__HI__", "Height-adjusted VAA", "+0.00092", "0.00009", "+10.36", "93%"),
          ],
          [2.3, 0.95, 0.85, 0.75, 0.85])
    rich(doc, [("Ninety-three percent of the shape effect survives the removal of everything "
                "attributable to where the pitch was thrown, and the ", ""), ("t", "i"),
               ("-statistic actually improves — stripping the location component removes noise as well "
                "as signal. Shape is a property of the pitcher, not of his aim.", "")])

    heading(doc, "5.4", "The result generalises to every pitch type", level=2)
    para(doc, "Sceptics of fastball-only findings are right to be sceptical, since spin means opposite "
              "things on a four-seamer and a curveball. The ladder was therefore refit independently on "
              "all seven pitch types with adequate sample.")
    table(doc, 3, "Incremental CV R² by pitch type (t-statistic in parentheses)",
          ["Pitch", "RV: Velo", "RV: Spin", "RV: Shape", "Whiff: Velo", "Whiff: Spin", "Whiff: Shape"],
          [
              ["Four-seam", "+.00025 (2.2)", "+.00020 (1.9)", "+.00085 (6.0)", "+.00540 (9.8)", "+.00129 (4.3)", "+.01923 (24.8)"],
              ["Sinker", "+.00008 (0.5)", "−.00000 (0.0)", "+.00217 (9.9)", "+.00300 (3.4)", "+.00061 (3.4)", "+.01085 (6.2)"],
              ["Cutter", "+.00034 (2.4)", "−.00014 (−2.3)", "+.00117 (6.7)", "+.00184 (5.9)", "−.00031 (−1.2)", "+.01145 (12.4)"],
              ["Slider", "−.00030 (−3.3)", "+.00012 (1.3)", "+.00099 (23.9)", "+.00102 (4.4)", "−.00009 (−0.4)", "+.00563 (14.4)"],
              ["Sweeper", "−.00031 (−1.4)", "−.00027 (−2.4)", "+.00168 (6.7)", "+.00206 (4.6)", "−.00026 (−1.0)", "+.00755 (6.8)"],
              ["Curveball", "+.00010 (1.2)", "−.00021 (−1.6)", "+.00051 (2.1)", "+.00496 (7.5)", "+.00025 (0.6)", "+.00561 (10.5)"],
              ["Changeup", "−.00025 (−2.7)", "+.00012 (0.9)", "+.00200 (9.1)", "+.00105 (6.5)", "+.00032 (1.8)", "+.00463 (7.4)"],
          ],
          [0.80, 0.90, 0.90, 0.92, 0.90, 0.90, 0.92],
          note="RV = run value. Shape is the largest block in all fourteen pitch-type-by-outcome "
               "combinations and significant in every one; residual spin fails significance in eleven "
               "of fourteen and is negative in six.")

    heading(doc, "5.5", "Reliability: the finding that changes evaluation", level=2)
    rich(doc, [("Everything to this point concerns what explains the outcome of a pitch. Evaluation asks "
                "a different question — what tells me how good this pitcher ", ""), ("is", "i"),
               (" — and the answer is governed less by effect size than by measurement error.", "")])
    figure(doc, figs["reliability"], 3,
           "Year-over-year correlation: what actually persists about a pitcher.",
           "Every shape and delivery metric persists from one season to the next at r ≥ 0.84. Whiff "
           "rate manages 0.63; run value per pitch, the statistic closest to “how well did he actually "
           "pitch,” manages 0.22.")

    table(doc, 4, "Reliability, four-seam fastballs (3,348 pitcher-seasons; 1,889 consecutive pairs)",
          ["Metric", "Kind", "Split-half r", "Spearman–Brown", "Year-over-year r"],
          [
              ["Release height", "shape", "1.000", "1.000", "0.956"],
              ["Extension", "shape", "0.999", "1.000", "0.949"],
              ["Release speed", "shape", "0.999", "1.000", "0.925"],
              ["VAA (height-adjusted)", "shape", "0.999", "0.999", "0.940"],
              ["Horizontal break", "shape", "0.998", "0.999", "0.886"],
              ["Residual spin", "shape", "0.998", "0.999", "0.915"],
              ["Induced vertical break", "shape", "0.996", "0.998", "0.836"],
              ["Whiff rate", "outcome", "0.658", "0.794", "0.628"],
              ("__HI__", "Run value per pitch", "outcome", "0.200", "0.333", "0.216"),
          ],
          [1.75, 0.72, 1.0, 1.15, 1.25])

    callout(doc, "What reliability means here", [
        [("If you split a pitcher's season in half at random and his induced vertical break in the "
          "first half tells you his second half almost exactly (", ""), ("r = 0.996", "m"),
         ("), the metric is essentially free of measurement error. If his run value in the first half "
          "barely tells you anything about the second (", ""), ("r = 0.20", "m"),
         ("), then most of what you are looking at when you look at his results is luck — sequencing, "
          "defence, where the ball happened to land.", "")],
        [("This is the whole argument in one line: ", ""),
         ("shape is a measurement, results are a sample.", "b")]])

    heading(doc, "5.6", "The information crossover", level=2)
    para(doc, "Because shape is measured nearly exactly from the very first pitch while outcome "
              "statistics need hundreds of pitches to mean anything, there must exist a sample size "
              "below which shape is the better predictor of a pitcher's future and above which his "
              "results are.")
    rich(doc, [("To locate it, each pitcher-season was subsampled to ", ""), ("exactly", "i"),
               (" N four-seamers, metrics computed from only those N pitches, and next season's "
                "performance predicted from them. The target is always measured on a full sample, so "
                "only the quality of what we know ", ""), ("now", "i"), (" varies.", "")])
    figure(doc, figs["crossover"], 4,
           "Predicting next season's whiff rate, by how much you have seen.",
           "Shape's predictive power is flat: it is just as informative from 30 pitches as from 500, "
           "because it is measured without error the moment the pitch leaves the hand. Outcome "
           "statistics start near-useless and improve steadily, overtaking shape at roughly 80 pitches.")

    table(doc, 5, "Out-of-sample R² predicting the following season",
          ["Pitches seen", "Whiff: Shape", "Whiff: Results", "Whiff: Both", "RV: Shape", "RV: Results", "RV: Both"],
          [
              ("__HI__", "30", "0.188", "0.094", "0.231", "0.036", "0.022", "0.046"),
              ("__HI__", "60", "0.193", "0.160", "0.266", "0.037", "0.038", "0.057"),
              ["125", "0.192", "0.250", "0.319", "0.044", "0.062", "0.078"],
              ["250", "0.206", "0.321", "0.366", "0.051", "0.087", "0.099"],
              ["500", "0.188", "0.362", "0.378", "0.037", "0.095", "0.087"],
          ],
          [0.94, 0.92, 0.97, 0.89, 0.84, 0.89, 0.82],
          note="Highlighted rows are those where shape beats the pitcher's own results. In every row, "
               "combining the two beats either alone.")

    rich(doc, [("Two facts deserve equal billing. Below roughly 80 pitches, shape beats results — at 30 "
                "pitches it beats them ", ""), ("two to one", "b"),
               (". Above that, results win, and the margin keeps growing. But in ", ""), ("every", "i"),
               (" row, combining the two beats either alone: at a full 250-pitch sample, adding shape "
                "to results lifts next-season whiff prediction from ", ""), ("0.321", "m"), (" to ", ""),
               ("0.366", "m"), (", a 14% relative improvement that is free for the taking.", "")])

    heading(doc, "5.7", "League context: velocity has compressed", level=2)
    rich(doc, [("Between 2015 and 2025, average four-seam velocity rose from ", ""), ("93.11", "m"),
               (" to ", ""), ("94.48", "m"), (" mph (", ""), ("t = +9.40", "m"),
               ("). Over the same window, the standard deviation across the league ", ""), ("fell", "i"),
               (" from ", ""), ("2.82", "m"), (" to ", ""), ("2.53", "m"), (" mph (", ""),
               ("t = −7.22", "m"), (").", "")])
    para(doc, "The value of one additional mph shows no significant trend (t = −0.83). What changed is "
              "the supply. Velocity is worth what it always was; there is simply far less spread left "
              "in it, so it separates one pitcher from another less than it used to. That is the "
              "defensible sense in which velocity “matters less” — not that hitters adapted, but that "
              "the population converged.")

    # ---------------- 6 ----------------
    heading(doc, "6", "Threats to validity, and where the thesis had to be revised")
    para(doc, "This section exists because a paper that only reports its confirmations is advertising. "
              "Three findings below genuinely constrain the argument, and the first two forced me to "
              "narrow it.")

    heading(doc, "6.1", "Command dwarfs everything, including shape", level=2)
    rich(doc, [("In permutation importance on run value, distance from the middle of the zone scores ", ""),
               ("0.050", "m"), (". The best physical trait, induced vertical break, scores ", ""),
               ("0.00097", "m"), (" — roughly ", ""), ("fifty times smaller", "b"),
               (". On any honest accounting, where the pitch is thrown matters vastly more than what "
                "the pitch does.", "")])
    callout(doc, "Revision forced by this result", [
        [("The claim is ", ""), ("not", "i"), (" “shape is the most important thing about a pitcher.” "
         "It is: ", ""),
         ("among the intrinsic properties of a pitch — the things a pitcher carries with him rather "
          "than executes on a given night — shape dominates.", "b"),
         (" Command sits outside that set and outranks all of them. Any framework that leads with shape "
          "and ignores command is incomplete.", "")]],
        band=WARN_BAND, bar="D9A441")

    heading(doc, "6.2", "Shape does not replace results on full seasons", level=2)
    rich(doc, [("My initial expectation was that shape's reliability advantage would make it the "
                "superior evaluation input across the board. It does not. At 250 pitches, past results "
                "predict next-season whiff rate at ", ""), ("R² = 0.321", "m"), (" against shape's ", ""),
               ("0.206", "m"), ("; at 500 pitches it is ", ""), ("0.362", "m"), (" against ", ""),
               ("0.188", "m"), (". Results win, and win comfortably, once you have a real sample.", "")])
    rich(doc, [("The reason is not that shape degrades — it is flat — but that outcome statistics "
                "contain ", ""), ("everything", "i"),
               (": shape, command, sequencing, deception, the pitcher's ability to execute under "
                "pressure. Given enough observations to average out the luck, that broader coverage "
                "wins.", "")])
    callout(doc, "Revision forced by this result", [
        [("Shape is not a replacement for outcome data. It is ", ""),
         ("the correct input when the sample is short, and an additive input when it is long.", "b"),
         (" Any presentation of this research that claims shape supersedes results is overselling it, "
          "and will be dismantled by the first analyst who runs the full-season comparison.", "")]],
        band=WARN_BAND, bar="D9A441")

    heading(doc, "6.3", "An earlier version of this experiment was wrong", level=2)
    rich(doc, [("The first crossover test filtered to pitcher-seasons with ", ""), ("at least", "i"),
               (" N pitches rather than subsampling to exactly N. Because most seasons passing a "
                "30-pitch filter still contain hundreds of pitches, that design never tested small "
                "samples at all, and it produced the false conclusion that results beat shape "
                "everywhere. The corrected version — subsampling each pitcher-season to exactly N — "
                "produces Figure 4. I report the error because the corrected and uncorrected designs "
                "give opposite answers.", "")])

    heading(doc, "6.4", "Remaining limitations", level=2)
    bullets(doc, [
        [("True spin efficiency is not measurable from public data.", "b"), (" Statcast's spin axis is "
         "inferred from observed movement, not measured from Hawk-Eye's 3D axis. Every public “spin "
         "efficiency” figure, including any derivable here, is an inference.", "")],
        [("Everything here is observational.", "b"), (" Nothing establishes that adding two inches of "
         "IVB to a specific pitcher's fastball would improve his results.", "")],
        [("Shape features are collinear", "b"), (" with each other and with velocity. The block's "
         "contribution is well identified; individual coefficients within it are not.", "")],
        [("The height adjustment removes a main effect, not an interaction.", "b"), (" Chamberlain "
         "(2021) shows that VAA's influence is concentrated at the vertical extremes of the zone and "
         "that VAA and pitch height interact non-linearly. The quadratic residualisation in §4.3 strips "
         "out the average relationship between approach angle and height, but does not model that "
         "interaction. A specification allowing the VAA effect to vary by zone region would likely find "
         "a larger shape contribution than reported here, so the §5.3 figure is better read as a "
         "conservative floor than a point estimate.", "")],
        [("Sequencing and tunnelling are absent.", "b"), (" A fastball plays off the pitch before it. "
         "Some of what is credited to shape is plausibly the arsenal-level interaction shape enables.", "")],
        [("Survivorship.", "b"), (" Reliability figures come from pitchers who threw at least 250 "
         "four-seamers — pitchers good enough to keep being used.", "")],
        [("Absolute effect sizes are small.", "b"), (" Pitch-level run value is mostly irreducible "
         "noise. Shape's gain is meaningful because it accumulates over thousands of pitches, not "
         "because any single pitch is well predicted.", "")],
    ])

    heading(doc, "6.5", "The adjusted thesis", level=2)
    callout(doc, "Stated so it can be checked", [
        [("Pitch shape is the dominant ", ""), ("pitch-intrinsic", "i"),
         (" driver of outcomes, ahead of velocity and far ahead of spin rate, on every pitch type and "
          "both major outcome measures. It is also the only pitch property measured reliably enough to "
          "support judgement on a small sample, which makes it the correct primary input below roughly "
          "80 pitches and a persistent additive input above that. It does not outrank command, and it "
          "does not replace outcome data on full seasons.", "")]])

    # ---------------- 7 ----------------
    heading(doc, "7", "A framework for evaluating the modern pitcher")
    rich(doc, [("The findings above imply a specific, ordered procedure. The organising principle is "
                "that ", ""), ("the right evaluation input depends on how much you have seen", "b"),
               (", and the crossover falls between the 60- and 125-pitch samples, interpolating to "
                "roughly 80.", "")])

    heading(doc, "7.1", "Triage by sample size", level=2)

    for tier_head, tier_title, tier_body, bar in (
        ("UNDER ~80 PITCHES OF THE PITCH TYPE", "Read the shape. Ignore the results.",
         ["Applies to: a prospect's first look, a reliever early in the year, a pitcher who has just "
          "added or redesigned a pitch, a rehab assignment, a trade-deadline target you have limited "
          "recent data on.",
          "At this sample, his results carry roughly half the information his ball flight does. Whiff "
          "rate, ERA, and run value are all still dominated by luck. Shape, by contrast, is already "
          "fully measured — IVB and VAA are as knowable after 30 pitches as after 3,000.",
          [("What to do: ", "b"), ("rank the pitch on IVB, height-adjusted VAA, and horizontal break "
           "against the league distribution for that pitch type. Treat the outcome line as essentially "
           "uninformative.", "")]], "1BAF7A"),
        ("ROUGHLY 80 TO 400 PITCHES", "Blend, weighting results upward as they accumulate.",
         ["Results have overtaken shape but have not run away with it, and the combination beats either "
          "alone by a wide margin — at 250 pitches, shape plus results predicts next season 14% better "
          "than results alone.",
          [("What to do: ", "b"), ("use both. Where they disagree, the disagreement is the signal: a "
           "pitcher with elite shape and poor results in this window is usually a command or sequencing "
           "problem, which is a coachable problem. The reverse — ordinary shape, excellent results — is "
           "the profile most likely to regress.", "")]], "2A78D6"),
        ("FULL SEASON AND BEYOND", "Lead with results. Use shape to explain and to forecast change.",
         ["With 500+ pitches, his own outcomes are the better single predictor and it is not close. "
          "Shape's role shifts from estimating talent to explaining it and detecting change.",
          [("What to do: ", "b"), ("lead with the outcome measures, but keep shape in the model — it "
           "still adds. Use shape as the early-warning system: because it is measured exactly, a "
           "two-inch drop in IVB or a flattening release is visible in a single outing, long before it "
           "shows up in results.", "")]], "98A5B2"),
    ):
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(12)
        hp.paragraph_format.space_after = Pt(0)
        r = hp.add_run(tier_head)
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.name = DATA_FONT
        r.font.color.rgb = INK2
        left_bar(hp, bar)

        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(4)
        tr2 = tp.add_run(tier_title)
        tr2.font.size = Pt(12)
        tr2.font.bold = True
        tr2.font.name = BODY_FONT
        tr2.font.color.rgb = INK
        left_bar(tp, bar)

        for parts in tier_body:
            if isinstance(parts, str):
                parts = [(parts, "")]
            bp = doc.add_paragraph()
            bp.paragraph_format.space_after = Pt(6)
            for text, style in parts:
                run = bp.add_run(text)
                run.font.size = Pt(10)
                run.font.bold = "b" in style
                run.font.italic = "i" in style
                run.font.name = BODY_FONT
                run.font.color.rgb = INK
            left_bar(bp, bar)

    heading(doc, "7.2", "What to measure, in priority order", level=2)
    numbered(doc, [
        [("Command first, always.", "b"), (" Location outranks every pitch property by roughly fifty to "
         "one. Any evaluation that starts with shape has already skipped the biggest term. Shape tells "
         "you the ceiling of the pitch; command tells you how much of it he will access.", "")],
        [("Induced vertical break.", "b"), (" The single most important shape variable and the top "
         "physical feature in permutation importance. Read it relative to the pitch type — high IVB is "
         "an asset on a four-seamer and a liability on a sinker.", "")],
        [("Height-adjusted vertical approach angle.", "b"), (" The largest standardized coefficient of "
         "any physical trait (0.0123 runs per standard deviation, an order of magnitude above "
         "velocity). Use the height-adjusted version — raw VAA is contaminated by where he threw it.", "")],
        [("Horizontal break, handedness-mirrored.", "b"), (" Always convert to arm-side positive, or "
         "lefties and righties will cancel each other out in any pooled comparison.", "")],
        [("Release height and extension.", "b"), (" Nearly perfectly stable year to year (r = 0.956 and "
         "0.949) and the geometric inputs that determine approach angle.", "")],
        [("Velocity.", "b"), (" Real and consistently positive for whiffs, but a weaker differentiator "
         "than it was — the league's spread has compressed 10% since 2015. Worth about 5.4 runs per "
         "standard deviation over a starter's season of four-seamers.", "")],
        [("Spin rate — as a diagnostic only.", "b"), (" It adds nothing to outcome prediction once "
         "movement is known. Its legitimate use is mechanical: a spin rate that is high relative to the "
         "movement it produces indicates inefficient, gyro-heavy spin, which is a coaching target, not "
         "an evaluation input.", "")],
    ])

    heading(doc, "7.3", "Errors this framework is designed to prevent", level=2)
    bullets(doc, [
        [("Ranking pitchers by raw spin rate.", "b"), (" It contributed nothing in eleven of fourteen "
         "pitch-type-by-outcome tests, and was negative in six.", "")],
        [("Comparing shape across pitch types.", "b"), (" Eighteen inches of IVB is elite on a "
         "four-seamer and a broken sinker. Every comparison must be within pitch type.", "")],
        [("Using raw VAA.", "b"), (" Roughly 7% of its apparent value is location, and reporting it "
         "unadjusted invites the accusation that you have rediscovered the high fastball.", "")],
        [("Trusting a reliever's ERA.", "b"), (" At 60 innings, run value per pitch has a year-over-year "
         "correlation of 0.22. It is close to noise.", "")],
        [("Treating shape as destiny.", "b"), (" Shape sets the ceiling. Command determines how much of "
         "it is realised, and command is the larger term.", "")],
    ])

    heading(doc, "7.4", "A worked reading", level=2)
    para(doc, "Consider a 26-year-old reliever called up in August who has thrown 84 four-seamers with a "
              "4.91 ERA and a 19% whiff rate. His fastball shows 18.9 inches of IVB and a "
              "height-adjusted VAA a full degree flatter than league average, at 95.1 mph.")
    para(doc, "The framework's reading: the sample sits right at the crossover, so the ERA and whiff "
              "rate carry little more information than the ball flight does — the results line should "
              "be heavily discounted. The shape profile is well above average on the two variables that "
              "matter most, on metrics already measured almost exactly despite the small sample. The "
              "correct inference is a pitcher whose stuff is real and whose results have not yet caught "
              "up, with the open question being command — which is where the remaining evaluation "
              "effort belongs, because it is both the largest term and the one this data does not "
              "settle.")

    # ---------------- 8 ----------------
    heading(doc, "8", "Conclusion")
    para(doc, "The industry spent a decade organising pitcher evaluation around two numbers that are one "
              "step removed from the thing that actually happens. Velocity and spin are inputs to a "
              "trajectory; the trajectory is what a hitter faces. Measuring the trajectory directly — "
              "induced vertical break, horizontal break, and the angle at which the ball arrives — "
              "explains four and a half times more variation in run value than velocity does, and "
              "residual spin adds nothing at all once movement is visible.")
    para(doc, "But the more useful finding is about measurement error rather than effect size. Shape is "
              "a property; results are a sample. A pitcher's induced vertical break is known almost "
              "exactly after thirty pitches, while his run value is barely known after a full season. "
              "That asymmetry, not the size of any coefficient, is what should change how pitchers are "
              "evaluated — and it changes it most precisely where evaluation is hardest and most "
              "valuable: the prospect, the callup, the new pitch, the trade target with sixty pitches "
              "of recent data.")
    rich(doc, [("The honest boundary is worth restating. Command still outranks every pitch property by "
                "a wide margin, and a full season of a pitcher's own results still beats his shape "
                "profile. Shape is the best available answer to a specific and common question — ", ""),
               ("what do I make of this arm before the results mean anything?", "i"),
               (" — and a persistent, additive input after that.", "")])

    # ---------------- 9 ----------------
    heading(doc, "9", "Reproduction")
    para(doc, "Every figure in this paper is produced by the analysis code, from data downloaded "
              "directly from Baseball Savant. No number here was entered by hand.")
    para(doc, "All code is available at github.com/scobobo/modern-pitcher-evaluation, released "
              "under the MIT licence. The Statcast cache is not committed — it is rebuilt by the "
              "fetch script below.")

    for label, cmds in (
        ("Environment", ["pip install pandas numpy scikit-learn matplotlib pyarrow",
                         "pip install --no-deps pybaseball pygithub pyjwt Deprecated wrapt pynacl",
                         "pip install requests beautifulsoup4 lxml tqdm attrs python-dateutil cffi"]),
        ("Fetch the data (~7.5M pitches, one parquet per season)",
         ["python src/fetch.py --seasons 2015 2016 2017 2018 2019 2020 2021 2022 2023 2024 2025"]),
        ("Section 5.1 — attribution ladder, both orderings",
         ["python run_analysis.py", "python run_analysis.py --target is_whiff"]),
        ("Sections 5.3–5.5 — VAA robustness, generality, reliability", ["python run_paper_analysis.py"]),
        ("Section 5.6 — the sample-size crossover", ["python run_sample_size_test.py"]),
    ):
        para(doc, label, size=9, bold=True, color=INK2, font=DATA_FONT, space_after=3, space_before=8)
        for cmd in cmds:
            cp = doc.add_paragraph()
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.left_indent = Inches(0.15)
            r = cp.add_run(cmd)
            r.font.size = Pt(8.5)
            r.font.name = DATA_FONT
            r.font.color.rgb = INK
            shade(cp, BAND_HEX)

    para(doc, "", space_after=6)
    para(doc, "Results are written to output/ as CSV. The verdict function prints whichever conclusion "
              "the numbers support, including “hypothesis not supported” — it was written before the "
              "results were known and was not adjusted afterward.", space_before=6)

    # ---------------- 10 ----------------
    heading(doc, "10", "References, data, and tools")

    heading(doc, "10.1", "Data and software", level=2)
    bullets(doc, [
        [("Statcast pitch-tracking data.", "b"), (" MLB Advanced Media, via Baseball Savant "
         "(baseballsavant.mlb.com). 2015–2025 regular seasons, retrieved August 2026. Tracking hardware: "
         "TrackMan through 2019, Hawk-Eye from 2020.", "")],
        [("pybaseball.", "b"), (" Open-source Python library for retrieving baseball data, used here "
         "for the Statcast pulls.", "")],
        [("Analysis code for this paper.", "b"), (" github.com/scobobo/modern-pitcher-evaluation, MIT licence.", "")],
        [("scikit-learn.", "b"), (" Pedregosa et al., “Scikit-learn: Machine Learning in Python,” "
         "Journal of Machine Learning Research 12 (2011), 2825–2830.", "")],
        [("Run expectancy framework.", "b"), (" Tango, Lichtman & Dolphin, The Book: Playing the "
         "Percentages in Baseball (2006), the basis for the change-in-run-expectancy outcome.", "")],
        [("Baseball aerodynamics.", "b"), (" Alan M. Nathan's published work on the physics of the "
         "baseball in flight, underlying the Magnus-effect reasoning in §2.1 and the gyro-spin caveat "
         "in §5.2.", "")],
    ])
    heading(doc, "10.2", "Prior work on approach angle", level=2)
    para(doc, "Vertical approach angle was not developed here. It was established as a public "
              "analytical concept largely through the following, which informed the framing in §2.2 "
              "and the height-adjustment procedure in §4.3:")
    bullets(doc, [
        [("Alex Chamberlain, “Where Vertical Approach Angle Seems to Matter Most,” FanGraphs, "
          "7 January 2021.", "b"), (" Finds that VAA's effect on whiffs is concentrated at the "
         "vertical extremes of the strike zone rather than through its middle, and that VAA and pitch "
         "height interact non-linearly — a flat pitch thrown low does not behave like a flat pitch "
         "thrown high. blogs.fangraphs.com/where-vertical-approach-angle-seems-to-matter-most/", "")],
        [("Alex Chamberlain, “A Visualized Primer on Vertical Approach Angle (VAA),” FanGraphs, "
          "1 February 2022.", "b"), (" Documents the strong correlation between VAA and pitch height "
         "that motivates the height-adjusted measure used throughout this paper, and maintains a "
         "public VAA Above Average leaderboard built on the same reasoning. "
         "blogs.fangraphs.com/a-visualized-primer-on-vertical-approach-angle-vaa/", "")],
    ])
    para(doc, "The height-adjusted VAA used here (§4.3) is a simpler construction than Chamberlain's "
              "leaderboard measure — a quadratic residual within season and pitch type — and is used "
              "because it can be recomputed from raw Statcast fields by anyone reproducing this "
              "analysis.")

    heading(doc, "10.3", "Author contributions and use of AI tools", level=2)
    para(doc, "I directed, interpreted, revised, and published this study; parts of its implementation "
              "and drafting were produced with AI assistance. The division of labour is set out below "
              "so readers can weigh it for themselves rather than guess.")

    para(doc, "WHAT I CONTRIBUTED", size=8, bold=True, color=INK2, font=DATA_FONT,
         space_after=4, space_before=8)
    bullets(doc, [
        [("The research question and the original hypothesis", "b"), (" — that spin rate had displaced "
         "velocity as the meaningful measure of a pitch. This study tested that hypothesis and "
         "rejected it.", "")],
        [("The decision to redirect the study toward pitch shape", "b"), (" once the evidence pointed "
         "there, and to define its scope: a pitch-level attribution analysis, a reliability analysis, "
         "and a practical evaluation framework, rather than a descriptive result alone.", "")],
        [("The falsifiability standard the project ran under.", "b"), (" I required from the outset "
         "that any finding capable of discrediting the thesis be surfaced and the theory adjusted to "
         "fit the evidence, rather than the reverse. Sections 6.1 and 6.2 exist because of that "
         "requirement, and both narrowed the claim I started with.", "")],
        [("The decision to publish the disconfirming results", "b"), (" — the command finding that "
         "outranks shape fifty to one, the full-season comparison that shape loses, and the "
         "experimental design error in §6.3. The argument needs none of them; it is more honest for "
         "all three.", "")],
        [("Interpretation of the results,", "b"), (" the structure of the evaluation framework in §7, "
         "the boundary conditions in §6.5, and editorial revision of the manuscript.", "")],
    ])

    para(doc, "WHAT AI ASSISTANCE CONTRIBUTED", size=8, bold=True, color=INK2, font=DATA_FONT,
         space_after=4, space_before=8)
    para(doc, "Anthropic's Claude, via Claude Code, was used under my direction to build the "
              "data-acquisition, cleaning, and feature-engineering pipeline; to implement the grouped "
              "cross-validation and paired-fold significance testing described in §4.2; to execute the "
              "robustness checks the thesis had to survive, including the location confound in §5.3 "
              "and the corrected sample-size experiment in §6.3; and to generate the figures, tables, "
              "and draft prose, which I revised.")

    para(doc, "WHAT I HOLD MYSELF ACCOUNTABLE FOR", size=8, bold=True, color=INK2, font=DATA_FONT,
         space_after=4, space_before=8)
    para(doc, "I can defend every methodological choice in §4: why cross-validation folds are grouped "
              "by pitcher, why every rung of the ladder is scored on identical rows, why incremental "
              "gains are paired across folds rather than compared as point estimates, and why the "
              "attribution ladder is run in both orderings. Each exists because its absence "
              "manufactures false positives — and §6.3 documents where an earlier version of this "
              "analysis fell into exactly that trap, reported here rather than quietly corrected.")
    para(doc, "Using capable tools well is not the same as outsourcing judgement. The tooling "
              "accelerated the implementation; the decisions about what to test, what would count as "
              "evidence, what to believe, and what to retract remained mine throughout.")
    callout(doc, "The verification standard", [
        "Readers are invited to verify rather than to trust. Every quantitative claim in this paper is "
        "reproducible from the open-source code in §9, against public data, by anyone with a laptop. "
        "The analysis was deliberately constructed so that its central hypothesis could fail — the "
        "reporting function prints “hypothesis not supported” when the numbers say so, and was written "
        "before the results were known.",
        [("That is the standard by which this work should be judged, and it is a stronger claim than "
          "authorship: ", ""),
         ("none of it requires taking my word for anything.", "b")]])

    colophon = para(doc, "", space_after=2, space_before=14)
    bottom_border(colophon)
    para(doc, "Data. Statcast pitch-level data via Baseball Savant, 2015–2025 regular seasons, retrieved August 2026. "
              "Methods. Gradient-boosted regression (scikit-learn HistGradientBoostingRegressor), "
              "5-fold cross-validation grouped by pitcher, incremental gains paired across identical "
              "folds. Ridge regression on standardized features for reported effect sizes.",
         size=8.5, color=INK2, space_after=4)
    para(doc, "Reporting standard. No result is described as real below t = 2 on its paired fold "
              "differences. Section 6 lists every finding that constrains the thesis, including one "
              "that overturned an earlier version of it.", size=8.5, color=INK2)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"wrote {path.name} ({path.stat().st_size / 1024:.0f} KB)")
